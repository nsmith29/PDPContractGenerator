#!/usr/bin/env python3

"""
    Author: Niamh Smith [niamh.smith.17@ucl.ac.uk]
    Date: 16/01/2025 [updated: 17/01/2025]


    Hey James!

    You'll want to run this in a commandline terminal like Terminal on your Mac using:
        python3 ./main.py

    You'll want to add in the details I've highlighted with comments '##' to make it work for you.

    You will also want to pip install numpy [pip3 install numpy] & docx [pip3 install python-docx] & datetime if you
    haven't already before running.

    If you have any questions or any issues with the code, please let me know :)
"""
import docx.document
import numpy as np
import datetime
from docx import Document

## List Contract types and file paths - name the variable the contract type and equal it to its template filepath.
contracts = {'NDA': "./PresentDayProduction Mutual NDA.docx"}



## Add the variable name of each contract type to list options.
options = ['NDA']



class bcolors:
    """
        Colours to be used for command line messages and for user inputs upon run of python file.

        Class definitions are each either a Select Graphic Rendition (SGR) or 8-bit Escape (ESC) color mode code.
    """

    KEYVAR = '\033[95m'  # was HEADER
    QUESTION = '\x1b[38;5;135m'
    CHOSEN = '\x1b[38;5;50m'
    INFO = '\x1b[38;5;221m'
    ACTION = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    ITALIC = '\033[3m'
    UNDERLINE = '\033[4m'



## Add to dictionary additional questions you'd like to be asked during creation of contract.
## I'm sure they'll be questions to ask which are specific to each contract type. Use key names for them in format of
##              'Q{contract type}{number of question for type}' i.e QNDA1
questions = {'Q1': f"{bcolors.QUESTION}What contract type would you like to create?{bcolors.ENDC}",
             'Q2': f"\n{bcolors.QUESTION}Is this contract with an individual?\n{bcolors.ACTION}Please answer "
                   f"{bcolors.CHOSEN}Y {bcolors.ACTION} or {bcolors.CHOSEN}N{bcolors.ENDC}",
             'Q3': f'\n{bcolors.QUESTION}What is the recipients name?{bcolors.ENDC}',
             'Q4': f"\n{bcolors.QUESTION}What is the recipients address?{bcolors.ACTION}"
                   f"\nPlease {bcolors.UNDERLINE}use '//' at the end of the address{bcolors.ENDC}{bcolors.ACTION} "
                   f"to signal you have finished inputting the address.{bcolors.ENDC}\n",
             'Q5': f"\n{bcolors.QUESTION}Which country is the recipient company in?{bcolors.ENDC}",
             'Q6': f"\n{bcolors.QUESTION}Who will be signing the contract on the recipient company's behalf?{bcolors.ENDC}",
             'QNDA1':f"\n{bcolors.QUESTION}What information does each party agree to disclose?{bcolors.ENDC}",
             'QNDA2':f"\n{bcolors.QUESTION}How long will this NDA be put in place for?{bcolors.ENDC}"}



class ErrMessages:
    @staticmethod
    def ValueErrorYorN():
        text = str(f"\n{bcolors.FAIL}WARNING:{bcolors.UNDERLINE}Invalid answer given!{bcolors.ENDC}{bcolors.INFO} "
                   f"\nOnly valid answers are {bcolors.CHOSEN}Y {bcolors.ACTION}and {bcolors.CHOSEN}N"
                   f"{bcolors.INFO}.{bcolors.ACTION} Try again. {bcolors.ENDC}")
        # pass text & lock to function to print each line of error slowly when lock is next unreleased & available.
        print(text)


    @staticmethod
    def ValueErrorlist():
        text = str(f"\n{bcolors.FAIL}ERROR: {bcolors.UNDERLINE}Contract type not implemented!{bcolors.ENDC}"
                   f"{bcolors.INFO} "
                   f"\nValid implemented contract types are:{bcolors.CHOSEN}{bcolors.ITALIC} \n"
                   +f"{', '.join(options)}"
                   +f"{bcolors.ENDC}")
        # pass text & lock to function to print each line of error slowly when lock is next unreleased & available.
        print(text)



def checkinput(func):
    """
         Testing for error in user's input response to questions asked.
    """
    def wrapper(Q: str, typ: str, ans: list=None, *args):
        """
            Inputs:
                Q(str)                  : Key for one of the questions within Core.PrePopDictsAndLists.questions.

                typ(str)                : Expected type of the answer and must match a key within func._format.

                ans(list)[optional]     : List of multiple valid answer options the user can choose from

                args                    : extra arguments such as a th.Lock or queue.Queue for controlling passing
                                          of into between threads and/or preventing simultaneous printing of
                                          multiple threads.

            Outputs:
                A(str or list)          : If error exception is triggered, return answer given by user when question
                                          asked again by function. If error exception is not triggered, return value
                                          of input A as given to function.
        """

        A = func(Q, typ, ans)
        if typ == 'list' or typ == 'YorN':
            try:
                if isinstance(A, list) == True:
                    if len([a for a in A if a not in ans]) != 0:
                        raise ValueError
                else:
                    if A not in ans:
                        raise ValueError
            except:
                # trigger error informing user that they've given an input which is invalid.
                eval("ErrMessages.ValueError{}()".format(typ))
                A = ask_question(Q, typ, ans)

        return A

    return wrapper



@checkinput
def ask_question(Q: str, typ: str, ans: list = None):
    """
        Asking required user input question.

        inputs:
            Q(str):               The key string within dictionary questions which corresponds to the question to be
                                  asked.

            typ(str):             Type of the user input expected as an answer for the question being asked. If
                                  expecting a single line of string as an answer use 'str', if expecting multiple lines
                                  of strings - such as an address being inputted - use 'multiple'. If expecting an
                                  answer of yes or no, use 'YorN'

            ans(list)[optional]:  List of expected answers, such as when the given user input must be within a certain
                                  selection of predetermined options. If there are no predetermined options for a
                                  question, you do not have to include this argument when calling ask_question.

        outputs:
            A(str or list):       A will usually be a string, only when typ of 'multiple' is given will A be a list.
    """
    _format = {'list': "A.split(', ') if ',' in A else [A]", 'YorN': 'A.upper()', 'str': 'A', 'multiple':'A'}
    assert Q in questions.keys(), "Q must be a key for one of the questions within ."
    assert typ in _format.keys(), "typ is the expected type of the answer and must match a key within _format."
    if ans:
        assert isinstance(ans, list), "ans must be a list of multiple valid answer options the user can choose from"
    print(questions[str(Q)])
    if typ != 'multiple':
        A = input()
        # if expected type is list, split string. If expected type is Yes or No, make sure string is upper case.
        A = eval(_format[typ])
    else:
        A, end = [], False
        while end == False:
            a = input()
            if '//' not in a:
                A.append(a)
            elif '//' in a:
                a = a.replace('//','')
                A.append(a)
                end = True

    return A



class program:
    """
        Main program which asks for user input and creates new contract docx from existing template given user input.
    """
    def __init__(self):
        # calculated the data at which script is being run.
        self.condate = f"{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}"

        print(f'{bcolors.KEYVAR} Welcome to PDP contract automation {bcolors.ENDC}\n ')

        # ask user for their input on questions by using the ask_question function
        self.contype = ask_question('Q1', 'list', options)

        self.individ_company = ask_question('Q2','YorN', ['Y','N'])

        self.repname = ask_question('Q3','str')

        self.repadd = ask_question('Q4','multiple')

        if self.individ_company == 'N':
            self.repcountry = ask_question('Q5', 'str')

            self.representative = ask_question('Q6','str')

        # specific contract type questions
        self.additional = [] # empty list to populate with the user input answers related to specific contract type
        for Q in [key for key in questions.keys() if self.contype[0] in key]:
            self.additional.append(ask_question(Q, 'str')) ## might need option for multiple lines to be taken in for question

        newfile = self.writingnew()
        print(f"{bcolors.CHOSEN}The file {newfile} has been made for {self.repname} from template for {self.contype[0]} contract{bcolors.ENDC}")


    def writingnew(self):
        document = Document(contracts[self.contype[0]])

        # record of lines in document as a list
        lines = self.updatelines(document)  # [x.text for x in document.paragraphs]

        # search array of lines for line index for full line containing date.
        idx4date = self.findIndex(lines, 'Date: ')
        # add actual date to the end of this line.
        document, lines = self.linereplace(document, " ".join([document.paragraphs[idx4date].text, self.condate]), idx4date)

        # exchange details in document that are specific for whether contract with an individual or a company.
        if self.individ_company == 'Y':
            document, lines = self.forindividual(lines, document)
        elif self.individ_company == 'N':
            document, lines = self.forcompany(lines, document)

        document, lines = eval("self.for{}(lines, document)".format(self.contype[0]))
        newfilename = "".join(['./',self.contype[0],'_for_', self.repname, '_',self.condate.replace('/','_'),'.docx'])
        document.save(newfilename)

        return newfilename


    def forNDA(self, lines: list[str], document: docx.document.Document):
        # get line & index of "insert detail ...", replace with additional[0] & replace doc line. Update lines.
        document, lines = self.linereplace(document, *self.repvarchange('[insert details e.g. discussing the possibility'
                                                     ' of the parties entering into a joint venture]', lines,
                                                                        self.additional[0]))

        # get line & index of "for number years", replace with additional[1] & replace doc line. Update lines.
        document, lines = self.linereplace(document, *self.repvarchange('[indefinitely][for [insert number]', lines,
                                                                       f'for {self.additional[1]}'))

        return document, lines

    def forindividual(self, lines: list[str], document: docx.document.Document):
        # get line replacement for name, index of line to replace & replace line in document. Update lines.
        document, lines = self.linereplace(document, *self.repnamechange('[NAME OF INDIVIDUAL]', lines))

        # get line replacement for address, index of line to replace & replace line in document.
        document, lines = self.linereplace(document, *self.repaddchange('[address of individual]', lines))

        # removal of lines which are associated with the contract being for a company
        document, lines = self.remove(document, 'OR', lines)
        document, lines = self.remove(document, '[NAME OF COMPANY]', lines)

        # signature setup at end of contract
        document, lines = self.linereplace(document, *self.PDPsignature(lines))
        document, lines = self.linereplace(document, *self.repsignature(lines, 'individual'))

        return document, lines

    def forcompany(self, lines: list[str], document: docx.document.Document):
        # get line replacement for name, index of line to replace & replace line in document. Update lines.
        document, lines = self.linereplace(document, *self.repnamechange('[NAME OF COMPANY]', lines))

        # get line replacement for address, index of line to replace & replace line in document. Update lines.
        document, lines = self.linereplace(document, *self.repaddchange('[ADDRESS]', lines))

        # get line replacement for '[COUNTRY]', index of line to replace & replace line in document. Update lines.
        document, lines = self.linereplace(document, *self.repvarchange('[COUNTRY]', lines, self.repcountry))

        # removal of lines which are associated with the contract being for an individual
        document, lines = self.remove(document, '[NAME OF INDIVIDUAL]', lines)
        document, lines = self.remove(document, 'OR', lines)

        # signature setup at end of contract
        document, lines = self.linereplace(document, *self.PDPsignature(lines))
        document, lines = self.linereplace(document, *self.repsignature(lines, 'company'))

        return document, lines

    def repsignature(self, lines, sign4):
        # dictionary of signing detail based on repicient type.
        signatures = {"individal": f"by {self.repname}",
                      "company": f"on behalf of {self.repname} by its duly authorised representative, {self.representative}"}

        return self.Signatures(lines, signatures[sign4])

    def PDPsignature(self, lines: list):

        return self.Signatures(lines, "by James Nugent, a representative of Cosmic Audio Ltd")

    def Signatures(self, lines: list, strg: str):
        # find line containing 'Signed [by' string
        signline, idx4sign = self.findlineandindex('Signed [by', lines)
        # replace all but the word 'Signed' in line with strg argument
        signline = self.repvarreplacement(signline, " ".join([x for x in signline.split() if x != "Signed"]), strg)

        return signline, idx4sign

    def linereplace(self, document: docx.document.Document, varline: str, idx4var: int):
        # replace line containing name in document
        document.paragraphs[idx4var].text = varline
        # update lines
        lines = self.updatelines(document)

        return document, lines


    def remove(self, document: docx.document.Document, strg: str, lines: list):
        # find line containing strg wanting to be removed
        line, inx2remove = self.findlineandindex(strg, lines)
        # clear the paragraph line within document
        document.paragraphs[inx2remove].clear()
        # update lines
        lines  = self.updatelines(document)

        return document, lines


    def updatelines(self, document: docx.document.Document):

        return [x.text for x in document.paragraphs]


    def repnamechange(self, strg: str, lines: list):

        return self.repvarchange(strg, lines, self.repname)


    def repvarchange(self, strg: str, lines: list, var: str):
        varlines, idx4lines = self.findlineandindex(strg, lines)
        varlines = self.repvarreplacement(varlines, strg, var)

        return varlines, idx4lines


    def repaddchange(self, strg: str, lines: list):
        # find line containing strg and index of line.
        addline, idx4add = self.findlineandindex(strg, lines)
        # replace strg within line with repadd.
        if len(self.repadd) > 1: # if user input for address was multiple lines
            if "," in self.repadd[0]:
                self.repadd = " ".join(self.repadd)# if commas already put at end of each address line
            else:
                self.repadd = ", ".join(self.repadd) # if commas not put at end of each address line
        else:
            self.repadd = self.repadd[0]
        addline = self.repvarreplacement(addline, strg, self.repadd)

        return addline, idx4add


    def repvarreplacement(self, varline: str, strg: str, var: str):

        # replace strg within line with var.
        return varline.replace(strg, var)


    def findlineandindex(self, strg: str, lines: list):
        # find total line containing strg in lines.
        line = [x for x in lines if strg in x][0]
        # identify index of this line.
        idx4add = self.findIndex(lines, line)

        return line, idx4add

    def findIndex(self, lines: list, line: str):

        return np.argwhere(np.asarray(lines) == line)[0][0]

if __name__ =='__main__':
    program()