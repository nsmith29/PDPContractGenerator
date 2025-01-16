from docx import Document

# Create a new Document object
document = Document()

# Add a heading
#document.add_heading('Contract', level=1)  # Level 1 is the highest heading level

# Add a paragraph
#paragraph = document.add_paragraph('This is a paragraph of text.')

# Save the document
document.save('contract.docx')